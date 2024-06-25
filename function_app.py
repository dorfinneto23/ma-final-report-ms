import azure.functions as func
import logging
import os #in order to get parameters values from azure function app enviroment vartiable - sql password for example 
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient # in order to use azure container storage
import io # in order to download pdf to memory and write into memory without disk permission needed 
import json # in order to use json 
from azure.servicebus import ServiceBusClient, ServiceBusMessage # in order to use azure service bus 
from openai import AzureOpenAI #for using openai services 
from azure.data.tables import TableServiceClient, TableClient, UpdateMode # in order to use azure storage table  
from azure.core.exceptions import ResourceExistsError, ResourceNotFoundError # in order to use azure storage table  exceptions 
import requests #in order to use translation function 
import uuid  #in order to use translation function 
import markdown2
from bs4 import BeautifulSoup
from docx import Document
import markdown
from docx.shared import Pt, RGBColor,Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile


# Azure Blob Storage connection string
connection_string_blob = os.environ.get('BlobStorageConnString')

#Azure service bus connection string 
connection_string_servicebus = os.environ.get('servicebusConnectionString')

#translate key
translate_key = os.environ.get('translate_key')



# Update field on specific entity/ row in storage table 
def update_cases_entity_field(table_name, partition_key, row_key, field_name, new_value,field_name2, new_value2):

    try:
        # Create a TableServiceClient using the connection string
        table_service_client = TableServiceClient.from_connection_string(conn_str=connection_string_blob)

        # Get a TableClient
        table_client = table_service_client.get_table_client(table_name)

        # Retrieve the entity
        entity = table_client.get_entity(partition_key, row_key)

        # Update the field
        entity[field_name] = new_value
        entity[field_name2] = new_value2

        # Update the entity in the table
        table_client.update_entity(entity, mode=UpdateMode.REPLACE)
        logging.info(f"update_cases_entity_field:Entity updated successfully.")

    except ResourceNotFoundError:
        logging.info(f"The entity with PartitionKey '{partition_key}' and RowKey '{row_key}' was not found.")
    except Exception as e:
        logging.info(f"An error occurred: {e}")

# Helper function to download blob content to stream 
def download_blob_stream(path):
        # Create a BlobServiceClient using the connection string
        container_name = "medicalanalysis"
        blob_service_client = BlobServiceClient.from_connection_string(connection_string_blob)
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(path)
        stream = io.BytesIO()
        blob_client.download_blob().download_to_stream(stream)
        stream.seek(0)
        return stream

#-------------------------------------------------------Markdown to DOCX Functions----------------------------------------
# Function to download image from Azure Blob Storage
def download_image_from_blob(image_url):
    response = requests.get(image_url)
    if response.status_code == 200:
        temp_image_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
        with open(temp_image_path, 'wb') as f:
            f.write(response.content)
        return temp_image_path
    else:
        raise Exception(f"Failed to download image: {response.status_code}")
    
# Function to add image to the header
def add_image_to_header(doc, image_path):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(2))  # Adjust width as needed
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def parse_html_to_docx(soup, document):

    # Download the image and get the local path
    image_path = download_image_from_blob("https://medicalanalysis.blob.core.windows.net/medicalanalysis/configuration/logo_doc.png")
    # Call the function to add image to the header
    add_image_to_header(document, image_path)
    
    def add_heading(text, level):
        heading = document.add_heading(level=level)
        run = heading.add_run(text)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Set the text direction to RTL
        set_rtl_direction(heading)

    def add_paragraph(text, bold=False):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Set the text direction to RTL
        set_rtl_direction(paragraph)

    for tag in soup.find_all(['h1', 'ol']):
        if tag.name == 'h1':
            add_heading(tag.text, level=1)
        elif tag.name == 'ol':
            for li in tag.find_all('li', recursive=False):
                strong_text = li.find('strong')
                if strong_text:
                    add_paragraph(strong_text.text, bold=True)
                ul = li.find('ul')
                if ul:
                    for ul_li in ul.find_all('li', recursive=False):
                        strong_text = ul_li.find('strong')
                        if strong_text:
                            add_paragraph(f"{strong_text.text} {ul_li.text.replace(strong_text.text, '').strip()}")

def set_rtl_direction(paragraph):
    """Sets the paragraph's text direction to RTL."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Add RTL property to paragraph
    paragraph_element = paragraph._element
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    paragraph_element.get_or_add_pPr().append(bidi)
    
    # Set the text direction for each run in the paragraph
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        bidi_run = OxmlElement('w:bidi')
        bidi_run.set(qn('w:val'), '1')
        rPr.append(bidi_run)

def set_docx_rtl(document):
    for paragraph in document.paragraphs:
        set_rtl_direction(paragraph)

def convert_txt_to_docx_with_reference(txt_blob_path, caseid,destination_folder):
    try:
        reference_docx_blob_path = "configuration/custom-reference.docx"

        # Download the markdown txt file content
        markdown_txt_stream = download_blob_stream(txt_blob_path)
        markdown_txt_content = markdown_txt_stream.getvalue().decode('utf-8')

        # Debug: Print markdown content
        logging.info(f"Markdown content: {markdown_txt_content}")

        # Convert Markdown content to HTML
        html_content = markdown2.markdown(markdown_txt_content)

        # Debug: Print HTML content
        logging.info(f"HTML content: {html_content}")

        # Parse HTML content
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Adjust HTML for RTL
        for tag in soup.find_all():
            tag['dir'] = 'rtl'
        html_content_rtl = str(soup)

        # Debug: Print adjusted HTML content
        logging.debug(f"RTL adjusted HTML content: {html_content_rtl}")
        
        # Save HTML content to a file
        html_file_name = "final_html.txt"
        save_final_files(html_content_rtl, caseid, html_file_name, destination_folder)
        
        doc = Document()

        # Set document direction to RTL
        set_docx_rtl(doc)

        # Add content to DOCX
        parse_html_to_docx(soup, doc)
        
        # Save the new DOCX document to a stream
        new_doc_stream = io.BytesIO()
        doc.save(new_doc_stream)
        new_doc_stream.seek(0)

        # Save the new DOCX document to Azure Storage
        doc_file_name = "final.docx"
        docx_path = save_final_files(new_doc_stream, caseid, doc_file_name, destination_folder)
        logging.info(f"Document saved to {docx_path}")
    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")


#-------------------------------------------------------Markdown to DOCX END ----------------------------------------

#Translate content given language 
def translate_text(text, to_language='he'):
    try:
        key = translate_key
        endpoint = "https://api.cognitive.microsofttranslator.com/"
        location = "eastus"
        path = '/translate'
        constructed_url = endpoint + path

        params = {
            'api-version': '3.0',
            'from': 'en',
            'to': [to_language]
        }

        headers = {
            'Ocp-Apim-Subscription-Key': key,
            'Ocp-Apim-Subscription-Region': location,
            'Content-type': 'application/json',
            'X-ClientTraceId': str(uuid.uuid4())
        }

        body = [{
            'text': text
        }]

        response = requests.post(constructed_url, params=params, headers=headers, json=body)
        response.raise_for_status()

        translations = response.json()
        translated_text = translations[0]['translations'][0]['text']
        return translated_text
    except Exception as e:
        logging.error(f"An error occurred:, {str(e)}")

#save files in "final" folder
def save_final_files(content,caseid,filename,file_folder):
    try:
        logging.info(f"save_ContentByClinicAreas start, content: {content},caseid: {caseid},filename: {filename}")
        container_name = "medicalanalysis"
        main_folder_name = "cases"
        folder_name="case-"+caseid
        blob_service_client = BlobServiceClient.from_connection_string(connection_string_blob)
        container_client = blob_service_client.get_container_client(container_name)
        basicPath = f"{main_folder_name}/{folder_name}"
        destinationPath = f"{basicPath}/final/{file_folder}/{filename}"
        # Upload the blob and overwrite if it already exists
        blob_client = container_client.upload_blob(name=destinationPath, data=content, overwrite=True)
        logging.info(f"the ContentByClinicAreas content file url is: {blob_client.url}")
        return destinationPath
    
    except Exception as e:
        logging.info(f"An error occurred:, {str(e)}")


# get content from txt file by path from azure table storage 
def get_content(path):
    try:
        logging.info(f"get_content function strating, path value: {path}")
        container_name = "medicalanalysis"
        blob_service_client = BlobServiceClient.from_connection_string(connection_string_blob)
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(path)
        download_stream = blob_client.download_blob()
        filecontent  = download_stream.read().decode('utf-8')
        logging.info(f"get_content: data from the txt file is {filecontent}")
        return filecontent
    except Exception as e:
        logging.error(f"get_content: Error update case: {str(e)}")
        return None    
    

#main function to create and manage all final filtered asistance response files 
def union_clinic_areas(table_name, caseid):

    destination_folder = "filtered"
    # Create a TableServiceClient object using the connection string
    service_client = TableServiceClient.from_connection_string(conn_str=connection_string_blob)
    
    # Get the table client
    table_client = service_client.get_table_client(table_name=table_name)
    

    # Query the table for entities with the given PartitionKey
    entities = table_client.query_entities(f"PartitionKey eq '{caseid}'")

    # union assistantResponsefiltered into one file for each entity
    combined_content = ""
    union_file_name = f"final-{caseid}.txt"
    for entity in entities:
        clinic_area = entity['RowKey']
        clinic_area_label = entity['clinicAreaLableName']
        content_path = entity['assistantResponsefiltered']
        filecontent = get_content(content_path)
        if filecontent!="":
            combined_content += "# " + clinic_area_label + "\n\n" + filecontent + "\n\n"
    #save union content of all clinic areas         
    save_final_files(combined_content,caseid,union_file_name,destination_folder)
    text_heb = translate_text(combined_content)
    heb_file_name = f"final-{caseid}-heb.txt"
    #save heb file
    heb_file_path = save_final_files(text_heb,caseid,heb_file_name,destination_folder)
    logging.info(f"union_clinic_areas: combined_content done")
    #convert heb txt file to docx 
    convert_txt_to_docx_with_reference(heb_file_path,caseid,destination_folder)
    
  #main function to create and manage all final filtered asistance response files 
def union_clinic_areas_disabilities_zero(table_name, caseid):

    destination_folder = "disabilities_zero"
    # Create a TableServiceClient object using the connection string
    service_client = TableServiceClient.from_connection_string(conn_str=connection_string_blob)
    
    # Get the table client
    table_client = service_client.get_table_client(table_name=table_name)
    

    # Query the table for entities with the given PartitionKey
    entities = table_client.query_entities(f"PartitionKey eq '{caseid}'")

    # union assistantResponsefiltered into one file for each entity
    combined_content = ""
    union_file_name = f"final-{caseid}-no-disabilities.txt"
    for entity in entities:
        clinic_area = entity['RowKey']
        clinic_area_label = entity['clinicAreaLableName']
        content_path = entity['assistantResponseNoDisabilities']
        filecontent = get_content(content_path)
        combined_content += "# " + clinic_area_label + "\n\n" + filecontent + "\n\n"
    #save union content of all clinic areas         
    save_final_files(combined_content,caseid,union_file_name,destination_folder)
    text_heb = translate_text(combined_content)
    heb_file_name = f"final-{caseid}-heb-no-disabilities.txt"
    #save heb file
    heb_file_path = save_final_files(text_heb,caseid,heb_file_name,destination_folder)
    logging.info(f"union_clinic_areas: combined_content done")

     

app = func.FunctionApp()

@app.service_bus_queue_trigger(arg_name="azservicebus", queue_name="final-report-process",
                               connection="medicalanalysis_SERVICEBUS") 
def finalReportMs(azservicebus: func.ServiceBusMessage):
    message_data = azservicebus.get_body().decode('utf-8')
    logging.info(f"Received messageesds: {message_data}")
    message_data_dict = json.loads(message_data)
    caseid = message_data_dict['caseid']
    #preparing final files where disabilities is not 0%
    union_clinic_areas_path = union_clinic_areas("contentByClinicAreas",caseid)
    logging.info(f"union_clinic_areas path: {union_clinic_areas_path}")
    #preparing final files where disabilities is  0%
    union_clinic_areas_path_disabilities_zero = union_clinic_areas_disabilities_zero("contentByClinicAreas",caseid)
    update_cases_entity_field("cases", caseid, "1", "status",13,"finalReportProcess",1)
   